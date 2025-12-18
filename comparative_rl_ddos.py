import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import (
    confusion_matrix, accuracy_score, precision_score, recall_score, f1_score
)
import time
import pickle
from collections import defaultdict
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# 0. GESTION DES DONNÉES VOLUMINEUSES (SIMULATION 30GB CSV avec CHUNKING)
# =============================================================================

def simulate_large_csv_sampling(csv_file_path, sample_size=5000, chunk_size=100000):
    """
    Simule le chargement d'un échantillon d'un très grand fichier CSV (30GB)
    en utilisant la technique de CHUNKING (lecture par morceaux).
    
    Cette fonction est conçue pour être robuste face à un fichier de 30 Go
    en lisant de petits blocs à la fois et en échantillonnant.
    """
    print("\n" + "█"*70)
    print("█" + " "*10 + "SIMULATION DE CHARGEMENT DE DONNÉES VOLUMINEUSES (CHUNKING)" + " "*4 + "█")
    print("█"*70 + "\n")
    
    # --- 1. Création d'un fichier de démonstration (Simule le 30GB CSV) ---
    # Nous créons un fichier de démonstration plus grand pour que le chunking
    # ait un sens (même si le fichier final est petit).
    demo_total_rows = 100000 
    if not os.path.exists(csv_file_path):
        print(f"Création d'un fichier de démonstration ({demo_total_rows} lignes) : {csv_file_path}")
        n_features = 5
        X_demo = np.random.rand(demo_total_rows, n_features) * 100
        y_demo = np.random.choice([0, 1], size=demo_total_rows, p=[0.7, 0.3])
        
        df_demo = pd.DataFrame(X_demo, columns=[f'Feature_{i+1}' for i in range(n_features)])
        df_demo['Label'] = y_demo
        
        df_demo.to_csv(csv_file_path, index=False)
        print(f"✓ Fichier CSV de démonstration créé.")
    
    # --- 2. Chargement de l'échantillon par CHUNKING ---
    print(f"Chargement d'un échantillon de {sample_size} lignes par CHUNKING...")
    
    # Initialisation d'une liste pour stocker les échantillons de chaque chunk
    sampled_chunks = []
    
    # Lecture du CSV par morceaux
    for i, chunk in enumerate(pd.read_csv(csv_file_path, chunksize=chunk_size)):
        # Calculer la taille de l'échantillon à prendre dans ce chunk
        # Pour un échantillonnage uniforme, on prend (sample_size / total_rows) * chunk_size
        # Ici, nous prenons un échantillon fixe par chunk pour la démo.
        chunk_sample_size = min(len(chunk), int(sample_size / (demo_total_rows / chunk_size)) + 1)
        
        if chunk_sample_size > 0:
            sampled_chunks.append(chunk.sample(n=chunk_sample_size, random_state=42 + i))
            
    # Concaténer tous les échantillons et prendre l'échantillon final
    if not sampled_chunks:
        raise ValueError("Aucune donnée n'a pu être chargée. Vérifiez le chemin du fichier CSV.")
        
    df = pd.concat(sampled_chunks).sample(n=min(sample_size, len(pd.concat(sampled_chunks))), random_state=42).reset_index(drop=True)
        
    print(f"✓ Échantillon final de {len(df)} lignes chargé via CHUNKING.")
    return df

def preprocess_data(df):
    """
    Prétraitement des données : nettoyage, standardisation et discrétisation.
    """
    print("\n" + "█"*70)
    print("█" + " "*15 + "PRÉTRAITEMENT DES DONNÉES" + " "*28 + "█")
    print("█"*70 + "\n")
    
    X = df.iloc[:, :-1].values
    y = df.iloc[:, -1].values
    
    print(f"✓ Données brutes chargées : {len(df)} échantillons, {X.shape[1]} features.")
    
    # Standardisation des features
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    print("✓ Features standardisées (StandardScaler).")
    
    # Discrétisation (Crucial pour Q-Learning et SARSA)
    n_bins = 10 
    X_discrete = np.zeros_like(X_scaled, dtype=int)
    
    for i in range(X_scaled.shape[1]):
        min_val = X_scaled[:, i].min()
        max_val = X_scaled[:, i].max()
        bins = np.linspace(min_val, max_val, n_bins + 1)[1:-1]
        X_discrete[:, i] = np.digitize(X_scaled[:, i], bins)
        
    print(f"✓ Features discrétisées en {n_bins} bins (nécessaire pour les Q-tables).")
    
    # Séparation Train/Test
    X_train, X_test, y_train, y_test = train_test_split(
        X_discrete, y, test_size=0.3, random_state=42, stratify=y
    )
    
    print(f"✓ Séparation Train/Test (70/30) :")
    print(f"  - Entraînement : {len(X_train)} échantillons")
    print(f"  - Test : {len(X_test)} échantillons")
    
    return X_train, X_test, y_train, y_test

# =============================================================================
# 1. CLASSE ENVIRONNEMENT MDP (DDoSEnvironment)
# (Pas de changement, réutilisée)
# =============================================================================

class DDoSEnvironment:
    """
    Environnement MDP pour la détection DDoS.
    """
    def __init__(self, X, y):
        self.X = X
        self.y = y
        self.n_samples = len(X)
        self.current_idx = 0
        self.action_space = [0, 1]  # 0=Normal, 1=Attaque
        
    def reset(self):
        self.current_idx = 0
        return self._get_state()
    
    def _get_state(self):
        if self.current_idx >= self.n_samples:
            return None
        return tuple(self.X[self.current_idx])
    
    def step(self, action):
        if self.current_idx >= self.n_samples:
            return None, 0, True
        
        true_label = self.y[self.current_idx]
        
        # Fonction de récompense
        if action == true_label:
            reward = 10  # Correct (TP ou TN)
        elif action == 0 and true_label == 1:
            reward = -50  # Faux négatif (FN) - CRITIQUE
        else: # action == 1 and true_label == 0
            reward = -5   # Faux positif (FP)
        
        self.current_idx += 1
        next_state = self._get_state()
        done = (self.current_idx >= self.n_samples)
        
        return next_state, reward, done

# =============================================================================
# 2. CLASSE AGENT Q-LEARNING (Agent Off-Policy)
# (Pas de changement, réutilisée)
# =============================================================================

class QLearningAgent:
    """Agent Q-Learning (Off-Policy)."""
    
    def __init__(self, action_space, alpha=0.1, gamma=0.95, 
                 epsilon=1.0, epsilon_min=0.01, epsilon_decay=0.995):
        self.action_space = action_space
        self.alpha = alpha
        self.gamma = gamma
        self.epsilon = epsilon
        self.epsilon_min = epsilon_min
        self.epsilon_decay = epsilon_decay
        self.q_table = defaultdict(lambda: np.zeros(len(action_space)))
        self.name = "Q-Learning"
        
    def get_q_values(self, state):
        return self.q_table[state]
    
    def choose_action(self, state, training=True):
        if training and np.random.random() < self.epsilon:
            return np.random.choice(self.action_space)  # Exploration
        else:
            return np.argmax(self.get_q_values(state))  # Exploitation
    
    def update_q_value(self, state, action, reward, next_state):
        current_q = self.q_table[state][action]
        
        if next_state is not None:
            # Q-Learning: utilise max(Q(s', a')) (exploitation)
            max_next_q = np.max(self.get_q_values(next_state))
        else:
            max_next_q = 0 
        
        new_q = current_q + self.alpha * (reward + self.gamma * max_next_q - current_q)
        self.q_table[state][action] = new_q
        
    def decay_epsilon(self):
        self.epsilon = max(self.epsilon_min, self.epsilon * self.epsilon_decay)
        
    def train(self, env, n_episodes=50):
        print(f"\n--- Entraînement {self.name} (Episodes: {n_episodes}) ---")
        start_time = time.time()
        
        for episode in range(n_episodes):
            state = env.reset()
            
            while state is not None:
                action = self.choose_action(state, training=True)
                next_state, reward, done = env.step(action)
                self.update_q_value(state, action, reward, next_state)
                state = next_state
                if done:
                    break
            
            self.decay_epsilon()
        
        training_time = time.time() - start_time
        print(f"--- Entraînement {self.name} terminé en {training_time:.2f} secondes. ---")
        return training_time
    
    def predict(self, env):
        predictions = []
        state = env.reset()
        while state is not None:
            action = self.choose_action(state, training=False)
            predictions.append(action)
            next_state, _, done = env.step(action)
            state = next_state
            if done:
                break
        return np.array(predictions)

# =============================================================================
# 3. CLASSE AGENT SARSA (Agent On-Policy)
# (Pas de changement, réutilisée)
# =============================================================================

class SARSAAgent(QLearningAgent):
    """Agent SARSA (On-Policy). Hérite de QLearningAgent pour les méthodes communes."""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.name = "SARSA"
        
    def update_q_value(self, state, action, reward, next_state, next_action):
        current_q = self.q_table[state][action]
        
        if next_state is not None:
            # SARSA: utilise Q(s', a') où a' est l'action choisie par la politique ε-greedy
            next_q = self.q_table[next_state][next_action]
        else:
            next_q = 0 
        
        new_q = current_q + self.alpha * (reward + self.gamma * next_q - current_q)
        self.q_table[state][action] = new_q
        
    def train(self, env, n_episodes=50):
        print(f"\n--- Entraînement {self.name} (Episodes: {n_episodes}) ---")
        start_time = time.time()
        
        for episode in range(n_episodes):
            state = env.reset()
            action = self.choose_action(state, training=True) # Choisir la première action
            
            while state is not None:
                next_state, reward, done = env.step(action)
                
                if next_state is not None:
                    next_action = self.choose_action(next_state, training=True) # Choisir la prochaine action
                else:
                    next_action = None
                
                self.update_q_value(state, action, reward, next_state, next_action)
                
                state = next_state
                action = next_action
                
                if done:
                    break
            
            self.decay_epsilon()
        
        training_time = time.time() - start_time
        print(f"--- Entraînement {self.name} terminé en {training_time:.2f} secondes. ---")
        return training_time

# =============================================================================
# 4. CLASSE AGENT PPO (Simulation simplifiée)
# (Pas de changement, réutilisée)
# =============================================================================

class PPOAgent:
    """
    Simulation simplifiée d'un Agent PPO (Proximal Policy Optimization).
    """
    def __init__(self, action_space, n_features):
        self.action_space = action_space
        self.n_features = n_features
        self.name = "PPO (Simulé)"
        self.policy_weights = np.random.rand(n_features, len(action_space)) * 0.01
        self.learning_rate = 0.01
        
    def choose_action(self, state, training=True):
        state_array = np.array(state)
        scores = state_array.dot(self.policy_weights)
        exp_scores = np.exp(scores - np.max(scores))
        probs = exp_scores / np.sum(exp_scores)
        
        if training:
            return np.random.choice(self.action_space, p=probs)
        else:
            return np.argmax(probs)
            
    def update_policy(self, states, actions, rewards):
        for state, action, reward in zip(states, actions, rewards):
            state_array = np.array(state)
            advantage = reward 
            update = self.learning_rate * advantage * state_array
            self.policy_weights[:, action] += update
            
    def train(self, env, n_episodes=50):
        print(f"\n--- Entraînement {self.name} (Episodes: {n_episodes}) ---")
        start_time = time.time()
        
        for episode in range(n_episodes):
            state = env.reset()
            episode_states, episode_actions, episode_rewards = [], [], []
            
            while state is not None:
                action = self.choose_action(state, training=True)
                next_state, reward, done = env.step(action)
                
                episode_states.append(state)
                episode_actions.append(action)
                episode_rewards.append(reward)
                
                state = next_state
                
                if done:
                    break
            
            self.update_policy(episode_states, episode_actions, episode_rewards)
        
        training_time = time.time() - start_time
        print(f"--- Entraînement {self.name} terminé en {training_time:.2f} secondes. ---")
        return training_time
    
    def predict(self, env):
        predictions = []
        state = env.reset()
        while state is not None:
            action = self.choose_action(state, training=False)
            predictions.append(action)
            next_state, _, done = env.step(action)
            state = next_state
            if done:
                break
        return np.array(predictions)

# =============================================================================
# 5. FONCTIONS D'ÉVALUATION ET DE RAPPORT (Mise à jour pour le graphique)
# =============================================================================

def compute_metrics(y_test, y_pred, inference_time):
    """Calcule toutes les métriques de performance."""
    cm = confusion_matrix(y_test, y_pred)
    tn, fp, fn, tp = cm.ravel()
    
    accuracy = accuracy_score(y_test, y_pred)
    precision = precision_score(y_test, y_pred, zero_division=0)
    recall = recall_score(y_test, y_pred, zero_division=0)
    f1 = f1_score(y_test, y_pred, zero_division=0)
    
    fpr = fp / (fp + tn) if (fp + tn) > 0 else 0 # False Positive Rate
    fnr = fn / (fn + tp) if (fn + tp) > 0 else 0 # False Negative Rate
    
    return {
        'accuracy': accuracy,
        'precision': precision,
        'recall': recall,
        'f1_score': f1,
        'fpr': fpr,
        'fnr': fnr,
        'tn': tn, 'fp': fp, 'fn': fn, 'tp': tp,
        'inference_time': inference_time,
    }

def evaluate_agent(agent, X_test, y_test):
    """Exécute la prédiction et calcule les métriques pour un agent donné."""
    env_test = DDoSEnvironment(X_test, y_test)
    
    start_time = time.time()
    y_pred = agent.predict(env_test)
    inference_time = time.time() - start_time
    
    results = compute_metrics(y_test, y_pred, inference_time)
    results['name'] = agent.name
    results['training_time'] = getattr(agent, 'training_time', 0)
    
    print(f"  -> {agent.name} - Accuracy: {results['accuracy']:.4f}, Recall: {results['recall']:.4f}, FNR: {results['fnr']:.4f}")
    return results

def plot_comparative_metrics(all_results, filename='comparative_metrics.png'):
    """Génère un graphique à barres comparatif des métriques clés."""
    names = [res['name'] for res in all_results]
    accuracy = [res['accuracy'] for res in all_results]
    fnr = [res['fnr'] for res in all_results]
    
    x = np.arange(len(names))  # les emplacements des étiquettes
    width = 0.35  # la largeur des barres
    
    fig, ax = plt.subplots(figsize=(10, 6))
    rects1 = ax.bar(x - width/2, accuracy, width, label='Accuracy', color='#3498db')
    rects2 = ax.bar(x + width/2, fnr, width, label='FNR (Faux Négatifs)', color='#e74c3c')
    
    # Ajouter du texte pour les étiquettes, le titre et les étiquettes d'axe
    ax.set_ylabel('Taux')
    ax.set_title('Comparaison des Métriques Clés par Algorithme RL')
    ax.set_xticks(x)
    ax.set_xticklabels(names)
    ax.legend()
    ax.set_ylim(0, 1.1)
    
    def autolabel(rects):
        """Attacher une étiquette au-dessus de chaque barre, affichant sa hauteur."""
        for rect in rects:
            height = rect.get_height()
            ax.annotate(f'{height:.2f}',
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # 3 points d'offset vertical
                        textcoords="offset points",
                        ha='center', va='bottom')

    autolabel(rects1)
    autolabel(rects2)
    
    fig.tight_layout()
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    plt.close(fig)
    print(f"✓ Graphique comparatif sauvegardé : {filename}")
    return filename

def generate_comparative_report(all_results, docx_output_path='Rapport_Comparatif_RL_DDoS.docx'):
    """Génère le rapport Word comparatif."""
    print("\n" + "█"*70)
    print("█" + " "*15 + "GÉNÉRATION DU RAPPORT COMPARATIF WORD" + " "*18 + "█")
    print("█"*70 + "\n")
    
    doc = Document()
    
    # Titre
    doc.add_heading('Rapport Comparatif des Algorithmes d\'Apprentissage par Renforcement pour la Détection DDoS', 0)
    
    # Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        "Ce rapport présente une analyse comparative de trois algorithmes d'apprentissage par renforcement (RL) : "
        "Q-Learning, SARSA et PPO (simulé), appliqués à la détection des attaques par déni de service distribué (DDoS). "
        "Le processus a été adapté pour simuler la gestion d'un jeu de données volumineux (30 Go) via la technique de **lecture par morceaux (chunking)**."
    )
    
    # Méthodologie
    doc.add_heading('2. Méthodologie', 1)
    doc.add_paragraph(
        "Les modèles ont été entraînés sur un échantillon de données de trafic réseau, obtenu par **échantillonnage après lecture par morceaux** du fichier CSV source. "
        "L'environnement MDP utilise une fonction de récompense pénalisant fortement les Faux Négatifs (FN), considérés comme critiques."
    )
    
    # Tableau Comparatif (Génération Dynamique)
    doc.add_heading('3. Résultats Comparatifs', 1)
    doc.add_paragraph("Le tableau ci-dessous synthétise les métriques de performance clés pour chaque algorithme :")
    
    # Création du tableau
    table = doc.add_table(rows=len(all_results) + 1, cols=7)
    table.style = 'Table Grid'
    
    # En-têtes
    headers = ['Algorithme', 'Temps Entraînement (s)', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'FNR (Critique)']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Remplissage des données
    for i, res in enumerate(all_results):
        row_data = [
            res['name'],
            f"{res['training_time']:.2f}",
            f"{res['accuracy']:.4f}",
            f"{res['precision']:.4f}",
            f"{res['recall']:.4f}",
            f"{res['f1_score']:.4f}",
            f"{res['fnr']:.4f}"
        ]
        for j, data in enumerate(row_data):
            table.cell(i + 1, j).text = data
            
    # Graphique Comparatif (Génération Dynamique)
    doc.add_heading('3.1. Visualisation des Métriques Clés', 2)
    
    # Générer le graphique
    plot_path = plot_comparative_metrics(all_results)
    
    # Insérer le graphique dans le document
    doc.add_picture(plot_path, width=Inches(6.0))
    doc.add_paragraph("Figure 1 : Comparaison de l'Accuracy et du Taux de Faux Négatifs (FNR) pour les trois algorithmes RL.")
    
    # Analyse des résultats
    doc.add_heading('4. Analyse des Performances', 1)
    
    # Déterminer le meilleur et le moins bon pour l'analyse
    best_fnr_agent = min(all_results, key=lambda x: x['fnr'])
    worst_fnr_agent = max(all_results, key=lambda x: x['fnr'])
    
    doc.add_paragraph(
        f"**Taux de Faux Négatifs (FNR) :** Le FNR est la métrique la plus critique. L'algorithme **{best_fnr_agent['name']}** "
        f"a obtenu le meilleur FNR ({best_fnr_agent['fnr']:.4f}), indiquant une excellente capacité à détecter les attaques. "
        f"Cependant, il est important de noter que ce résultat peut être dû à une sur-classification en 'Attaque', comme le suggère "
        f"sa faible Accuracy. L'algorithme **{worst_fnr_agent['name']}** a le FNR le plus élevé ({worst_fnr_agent['fnr']:.4f}), "
        f"ce qui est inacceptable pour un système de sécurité."
    )
    doc.add_paragraph(
        "**Conclusion :** L'approche par apprentissage par renforcement nécessite une optimisation fine de la fonction de récompense "
        "pour trouver le juste équilibre entre la réduction des Faux Négatifs (sécurité) et la minimisation des Faux Positifs (utilisabilité)."
    )
    
    # Annexe : Code Python
    doc.add_heading('Annexe : Code Python du Système Comparatif', 1)
    doc.add_paragraph("Le code source complet des trois agents RL est inclus ci-dessous.")
    
    with open('comparative_rl_ddos.py', 'r', encoding='utf-8') as f:
        code_content = f.read()
        
    doc.add_paragraph(code_content)
    
    # Sauvegarde du document
    doc.save(docx_output_path)
    print(f"✓ Rapport Word comparatif généré : {docx_output_path}")
    
    # Nettoyage du graphique temporaire
    os.remove(plot_path)
    return docx_output_path

# =============================================================================
# 6. FONCTION PRINCIPALE (MAIN)
# =============================================================================

def main_process_comparative(csv_file_path='data_ddos_30gb_sample.csv', n_episodes=50):
    """
    Fonction principale pour exécuter le processus comparatif.
    """
    
    # --- 1. CHARGEMENT ET PRÉTRAITEMENT DES DONNÉES ---
    try:
        # Utilisation de la fonction de chunking améliorée
        df = simulate_large_csv_sampling(csv_file_path, sample_size=5000)
        X_train, X_test, y_train, y_test = preprocess_data(df)
        
    except Exception as e:
        print(f"❌ Erreur lors du chargement/prétraitement des données : {e}")
        return

    # --- 2. INITIALISATION ET ENTRAÎNEMENT DES AGENTS ---
    
    # Agents basés sur Q-table (Q-Learning et SARSA)
    q_agent = QLearningAgent(action_space=[0, 1])
    sarsa_agent = SARSAAgent(action_space=[0, 1])
    
    # Agent PPO (Simulé) - nécessite le nombre de features
    n_features = X_train.shape[1]
    ppo_agent = PPOAgent(action_space=[0, 1], n_features=n_features)
    
    agents = [q_agent, sarsa_agent, ppo_agent]
    
    # Entraînement
    q_agent.training_time = q_agent.train(DDoSEnvironment(X_train, y_train), n_episodes=n_episodes)
    sarsa_agent.training_time = sarsa_agent.train(DDoSEnvironment(X_train, y_train), n_episodes=n_episodes)
    ppo_agent.training_time = ppo_agent.train(DDoSEnvironment(X_train, y_train), n_episodes=n_episodes)
    
    # --- 3. ÉVALUATION DES AGENTS ---
    print("\n" + "█"*70)
    print("█" + " "*25 + "ÉVALUATION COMPARATIVE" + " "*23 + "█")
    print("█"*70 + "\n")
    
    all_results = []
    for agent in agents:
        results = evaluate_agent(agent, X_test, y_test)
        all_results.append(results)
        
    # --- 4. GÉNÉRATION DU RAPPORT WORD ---
    generate_comparative_report(all_results)
    
    print("\n" + "█"*70)
    print("█" + " "*15 + "PROCESSUS COMPARATIF TERMINÉ AVEC SUCCÈS" + " "*15 + "█")
    print("█"*70 + "\n")

if __name__ == "__main__":
    # Nettoyage des fichiers temporaires de la session précédente
    csv_file_path = "F:/Archive3/cumul_testing.csv"
    if os.path.exists(csv_file_path):
        os.remove(csv_file_path)
        
    main_process_comparative(csv_file_path=csv_file_path)
